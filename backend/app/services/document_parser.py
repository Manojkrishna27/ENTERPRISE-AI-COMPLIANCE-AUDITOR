import os
import fitz  # PyMuPDF
import docx
from flask import current_app
from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter

class DocumentParser:
    @classmethod
    def parse_pdf(cls, file_path):
        """
        Parses a PDF using PyMuPDF (fitz) and extracts text blocks,
        preserving page numbers and paragraph (block) numbers.
        """
        chunks = []
        doc = fitz.open(file_path)
        
        chunk_position = 0
        for page_idx, page in enumerate(doc):
            page_num = page_idx + 1
            # "blocks" returns: (x0, y0, x1, y1, text, block_no, block_type)
            blocks = page.get_text("blocks")
            
            for block in blocks:
                # Replace non-breaking spaces and strip
                text = block[4].replace('\xa0', ' ').strip()
                block_no = block[5]
                block_type = block[6]
                
                # Filter out image blocks (block_type == 1) or very short/whitespace blocks
                if block_type == 0 and text:
                    chunks.append({
                        "text": text,
                        "page_number": page_num,
                        "paragraph_number": block_no + 1
                    })
                    
        doc.close()
        return cls._apply_semantic_chunking(chunks)

    @classmethod
    def parse_docx(cls, file_path):
        """
        Parses a DOCX using python-docx and extracts paragraphs and tables.
        Word files do not have rigid pages, so pages are simulated
        every 10 paragraphs/table rows.
        """
        chunks = []
        doc = docx.Document(file_path)
        
        chunk_position = 0
        element_idx = 0
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.replace('\xa0', ' ').strip()
            if text:
                simulated_page = (element_idx // 10) + 1
                chunks.append({
                    "text": text,
                    "page_number": simulated_page,
                    "paragraph_number": element_idx + 1
                })
                element_idx += 1
                
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.replace('\xa0', ' ').strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    text = " | ".join(row_data)
                    simulated_page = (element_idx // 10) + 1
                    chunks.append({
                        "text": f"[TABLE ROW]: {text}",
                        "page_number": simulated_page,
                        "paragraph_number": element_idx + 1
                    })
                    element_idx += 1
                
        return cls._apply_semantic_chunking(chunks)

    @classmethod
    def _apply_semantic_chunking(cls, raw_chunks):
        """
        Applies LlamaIndex SentenceSplitter (512 tokens, 64 overlap).
        Documents are created per paragraph to retain precise page/paragraph metadata.
        """
        docs = [
            Document(
                text=c["text"], 
                metadata={"page_number": c["page_number"], "paragraph_number": c["paragraph_number"]}
            ) 
            for c in raw_chunks
        ]
        
        splitter = SentenceSplitter(chunk_size=512, chunk_overlap=64)
        nodes = splitter.get_nodes_from_documents(docs)
        
        final_chunks = []
        for idx, node in enumerate(nodes):
            final_chunks.append({
                "text": node.text,
                "page_number": node.metadata.get("page_number", 1),
                "paragraph_number": node.metadata.get("paragraph_number", 1),
                "chunk_position": idx
            })
            
        return final_chunks

    @classmethod
    def parse_document(cls, file_path, file_type):
        """
        Routes the document to the correct parser based on file type.
        """
        file_type = file_type.upper()
        if file_type == 'PDF':
            return cls.parse_pdf(file_path)
        elif file_type == 'DOCX' or file_type == 'DOC':
            return cls.parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
current_parser = DocumentParser()
